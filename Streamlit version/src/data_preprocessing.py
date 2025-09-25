import pdfplumber
from docx import Document
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from config import Config

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeTextExtractor:
    """Extract and clean text from various resume formats"""
    
    def __init__(self, config: Optional[Config] = None):
        self.config = config or Config()
        
    def extract_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF resume"""
        try:
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text + "\n"
                    
                    # Extract table data if present
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += " | ".join([cell for cell in row if cell]) + "\n"
            
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting PDF {pdf_path}: {e}")
            return ""
    
    def extract_from_docx(self, docx_path: Path) -> str:
        """Extract text from DOCX resume"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX {docx_path}: {e}")
            return ""
    
    def extract_from_txt(self, txt_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting TXT {txt_path}: {e}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\-\(\)\[\]\/\@\+\#\&\%\$]', '', text)
        
        # Normalize common resume sections
        text = re.sub(r'(?i)(experience|education|skills|projects|certifications):', r'\n\1:\n', text)
        
        # Remove page markers
        text = re.sub(r'--- Page \d+ ---', '', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Optional: Remove email/phone for privacy
        # text = re.sub(r'\S+@\S+', '[EMAIL]', text)
        # text = re.sub(r'\b\d{3}-\d{3}-\d{4}\b', '[PHONE]', text)
        
        return text.strip()
    
    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information from resume text"""
        contact_info = {
            "emails": [],
            "phones": [],
            "linkedin": [],
            "websites": []
        }
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        contact_info["emails"] = re.findall(email_pattern, text)
        
        # Phone extraction
        phone_patterns = [
            r'\b\d{3}-\d{3}-\d{4}\b',  # 123-456-7890
            r'\b\(\d{3}\)\s*\d{3}-\d{4}\b',  # (123) 456-7890
            r'\b\d{10}\b',  # 1234567890
            r'\+\d{1,3}\s*\d{3,4}\s*\d{3,4}\s*\d{3,4}'  # International
        ]
        
        for pattern in phone_patterns:
            contact_info["phones"].extend(re.findall(pattern, text))
        
        # LinkedIn extraction
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        contact_info["linkedin"] = re.findall(linkedin_pattern, text.lower())
        
        # Website extraction
        website_pattern = r'https?://(?:[-\w.])+(?:[:\d]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:[\w.])*)?)?'
        contact_info["websites"] = re.findall(website_pattern, text)
        
        return contact_info
    
    def process_resume_folder(self, input_folder: Path, output_folder: Path) -> Dict[str, Any]:
        """Process all resumes in a folder and return processing stats"""
        input_folder = Path(input_folder)
        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)
        
        extractors = {
            '.pdf': self.extract_from_pdf,
            '.docx': self.extract_from_docx, 
            '.txt': self.extract_from_txt
        }
        
        stats = {
            "processed": 0,
            "failed": 0,
            "files": []
        }
        
        for resume_file in input_folder.iterdir():
            if resume_file.suffix.lower() in extractors:
                logger.info(f"Processing: {resume_file.name}")
                
                try:
                    # Extract text
                    extractor = extractors[resume_file.suffix.lower()]
                    text = extractor(resume_file)
                    
                    if text:
                        # Save processed text
                        output_file = output_folder / f"{resume_file.stem}.txt"
                        with open(output_file, 'w', encoding='utf-8') as f:
                            f.write(text)
                        
                        # Extract contact info
                        contact_info = self.extract_contact_info(text)
                        contact_file = output_folder / f"{resume_file.stem}_contact.json"
                        with open(contact_file, 'w', encoding='utf-8') as f:
                            json.dump(contact_info, f, indent=2)
                        
                        stats["processed"] += 1
                        stats["files"].append({
                            "original": resume_file.name,
                            "processed": output_file.name,
                            "contact": contact_file.name,
                            "text_length": len(text),
                            "status": "success"
                        })
                        
                        logger.info(f"Successfully processed: {resume_file.name}")
                    else:
                        stats["failed"] += 1
                        stats["files"].append({
                            "original": resume_file.name,
                            "status": "failed - no text extracted"
                        })
                        logger.warning(f"No text extracted from: {resume_file.name}")
                        
                except Exception as e:
                    stats["failed"] += 1
                    stats["files"].append({
                        "original": resume_file.name,
                        "status": f"failed - {str(e)}"
                    })
                    logger.error(f"Failed to process {resume_file.name}: {e}")
        
        return stats

class DataValidator:
    """Validate extracted resume data"""
    
    @staticmethod
    def validate_resume_text(text: str) -> Dict[str, Any]:
        """Validate if text looks like a resume"""
        if not text or len(text.strip()) < 100:
            return {"valid": False, "reason": "Text too short"}
        
        # Check for common resume keywords
        resume_keywords = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'degree', 'certification', 'project',
            'responsibility', 'achievement', 'accomplishment'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
        
        if keyword_count < 3:
            return {"valid": False, "reason": "Insufficient resume keywords"}
        
        # Check for personal information indicators
        has_contact = bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text))
        has_name = bool(re.search(r'^[A-Z][a-z]+ [A-Z][a-z]+', text.strip()))
        
        confidence = min((keyword_count / len(resume_keywords)) + (0.2 if has_contact else 0) + (0.1 if has_name else 0), 1.0)
        
        return {
            "valid": confidence > 0.3,
            "confidence": confidence,
            "keyword_count": keyword_count,
            "has_contact": has_contact,
            "has_name": has_name
        }

# Utility functions
def batch_process_resumes(input_dir: str, output_dir: str) -> Dict[str, Any]:
    """Batch process multiple resumes"""
    extractor = ResumeTextExtractor()
    return extractor.process_resume_folder(Path(input_dir), Path(output_dir))

if __name__ == "__main__":
    # Test the extraction
    Config.create_directories()
    
    # Example usage
    extractor = ResumeTextExtractor()
    
    # Process sample resume
    sample_text = """
    John Doe
    Software Engineer
    john.doe@email.com | (555) 123-4567
    
    EXPERIENCE:
    Senior Python Developer at Tech Corp (2020-2023)
    - Developed web applications using Django and React
    - Worked with PostgreSQL databases and AWS cloud services
    """
    
    cleaned_text = extractor.clean_text(sample_text)
    contact_info = extractor.extract_contact_info(sample_text)
    
    print("=== TEXT EXTRACTION TEST ===")
    print("Cleaned text:", cleaned_text[:200] + "...")
    print("Contact info:", json.dumps(contact_info, indent=2))
    
    # Validate
    validator = DataValidator()
    validation = validator.validate_resume_text(cleaned_text)
    print("Validation:", json.dumps(validation, indent=2))
